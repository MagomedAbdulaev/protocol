import os
from docx import Document
from django.http import FileResponse
from django.conf import settings
from .forms import *
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth.hashers import check_password
from django.contrib.auth import logout, login
from django.core.exceptions import ObjectDoesNotExist


@login_required
def home(request):
    user = request.user
    company = user.company

    context = {
        'title': 'Главная',
        'user': user,
        'company': company,
    }

    return render(request, 'home.html', context)

def login_user(request):
    companies = Company.objects.all()
    if request.method == 'POST':
        login_form = LoginUserForm(request.POST)
        if login_form.is_valid():
            try:
                user = Person.objects.get(username=request.POST['username'])
                if check_password(request.POST['password'], user.password) or request.POST['password'] == user.password:
                    try:
                        if user.company.name == request.POST['company']:
                            login(request, user)
                            return redirect('protocol_app:home')
                        else:
                            login_form.add_error('username', f'У пользователя {user.username} нет компании {request.POST["company"]}')
                    except AttributeError:
                        login_form.add_error('username', f'У пользователя {user.username} нет привязанных компаний')
                else:
                    login_form.add_error('password', 'Неверный пароль')
            except ObjectDoesNotExist:
                login_form.add_error('username', 'Пользователя с таким логином не существует')
    else:
        login_form = LoginUserForm()

    context = {
        'title': 'Авторизация',
        'login_form': login_form,
        'companies': companies,
        'header_is_none': True,
    }

    return render(request, 'login.html', context)

def logout_user(request):
    logout(request)
    return redirect('protocol_app:login_user')


@login_required
def upload_pdf(request):
    pdf_file_path = request.session.get('pdf_file', False)
    
    if pdf_file_path:
        filename = os.path.basename(pdf_file_path)
        response = FileResponse(open(pdf_file_path, 'rb'), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    else:
        return redirect('protocol_app:home')

@login_required
def create_user(request):
    if request.method == 'POST':
        create_employee_form = CreateEmployeeForm(request.POST)
        if create_employee_form.is_valid():
            instance = create_employee_form.save(commit=False)
            instance.company = request.user.company
            instance.save()

            document_path = os.path.join(settings.BASE_DIR, 'tz', 'шаблоны', 'С БДД.docx')

            if not os.path.exists(document_path):
                raise FileNotFoundError(f"The file at {document_path} was not found.")

            document = Document(document_path)

            replacements = {
                'Компания': request.user.company.name,
                '№': request.user.company.id,
                'Дата протокола': request.POST['date_check'],
                'Член1': request.POST['name_chairman'],
                'Д1': request.POST['post_chairman'],
                'Член2': request.POST['name_first_member_commission'],
                'Д2': request.POST['post_first_member_commission'],
                'Член3': request.POST['name_second_member_commission'],
                'Д3': request.POST['post_second_member_commission'],
                'номер программы': instance.special.id,
                'для кого': instance.special.post,
                'ФИО': request.POST['full_name'],
                'Должность': instance.special.post,
                'Причина': request.POST['reason'],
                'Стаж': request.POST['work_experience'],
                'группа': instance.special.group_eb,
                'Дата ЭБ': request.POST['date_check'],
                'Инструктаж': request.POST['fire_safety_instruction'],
                'Энф': request.POST['responsible_electrical_industry'],
            }

            def replace_text_in_paragraph(paragraph, replacements):
                for key, value in replacements.items():
                    if f'{{{key}}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value))

            def replace_text_in_runs(paragraph, replacements):
                for run in paragraph.runs:
                    for key, value in replacements.items():
                        if f'{{{key}}}' in run.text:
                            run.text = run.text.replace(f'{{{key}}}', str(value))

            def replace_text_in_table(table, replacements):
                for row in table.rows:
                    for cell in row.cells:
                        for par in cell.paragraphs:
                            replace_text_in_paragraph(par, replacements)
                            replace_text_in_runs(par, replacements)
                        for nested_table in cell.tables:
                            replace_text_in_table(nested_table, replacements)

            for par in document.paragraphs:
                replace_text_in_paragraph(par, replacements)
                replace_text_in_runs(par, replacements)

            for table in document.tables:
                replace_text_in_table(table, replacements)

            temp_docx_path = os.path.join(settings.MEDIA_ROOT, f'docx/{request.POST["full_name"]}.docx')
            document.save(temp_docx_path)

            # pdf_path = os.path.join(settings.MEDIA_ROOT, f'pdf/{request.POST["full_name"]}.pdf')
            # request.session['pdf_file'] = pdf_path
            request.session.modified = True
            
            # Convert DOCX to PDF

            return redirect('protocol_app:upload_pdf')

    else:
        create_employee_form = CreateEmployeeForm()

    context = {
        'title': f'Форма для создания удостоверений в компанию {request.user.company.name}',
        'header_is_none': True,
        'create_employee_form': create_employee_form,
    }

    return render(request, 'create_user.html', context)
    
    